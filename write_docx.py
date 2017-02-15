from  docx import Document

document=Document()


document.add_heading('Heading, level ', level=0)
# document.add_heading('Heading, level ', level=1)
# document.add_heading('Heading, level ', level=3)
# document.add_heading('Heading, level ', level=2)
# document.add_heading('Heading, level ', level=4)
# document.add_heading('Heading, level ', level=5)
# document.add_heading('Heading, level ', level=6)
p=document.add_paragraph('我是段落')
p.add_run('我要加粗').bold= True
p.add_run('and some')
p.add_run('我是斜体').italic=True

document.add_paragraph('Intense quote',style='IntenseQuote')
document.add_paragraph('我前面有标点符号',style='ListBullet')
document.add_paragraph('我前面是数字',style='ListNumber')
# document.add_picture('test.png',width=Inches(1.25))
table=document.add_table(rows=1,cols=3)
header_cells=table.rows[0].cells
#header_cells[0].text='Qty'
#table[0][0]='Qty' #error
header_cells[0].text='Qty'
header_cells[1].text='Id'
header_cells[2].text='Desc'
document.add_page_break()
document.save('test.docx')










# def readdocx(docname):
# 	fulltext=[]
# 	doc=docx.Document(docname)
# 	paras = doc.paragraphs
# 	for x in  paras:
# 		fulltext.append(x.text)
# 	return  '\n'.join(fulltext) #原本换行作为分割标志，
# 	#print将每一段落用''框起来；用了这个则用换行表示出段落
# 	# return  fulltext
# print(readdocx('用户访谈.docx'))
